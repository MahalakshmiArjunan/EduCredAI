"""Downloadable document builders for chapter key points and question papers."""
from io import BytesIO
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


BRAND_BLUE = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x12, 0x1C, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def _header(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Project Vidya")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_BLUE

    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.bold = True
    hr.font.size = Pt(22)
    hr.font.color.rgb = DARK

    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED
    doc.add_paragraph()  # spacer


def _bullets_from_chunk(chunk: str) -> list:
    """Split a content chunk sentence-by-sentence into bullet points."""
    if not chunk:
        return []
    # Simple sentence split, keeping short chunks together
    import re
    parts = re.split(r"(?<=[.!?])\s+", chunk.strip())
    return [p.strip() for p in parts if len(p.strip()) > 4]


def build_key_points_docx(chapter: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _header(
        doc,
        title=f"Key Points · {chapter.get('title','Chapter')}",
        subtitle=(
            f"Grade {chapter.get('grade','')} · {chapter.get('subject','')} · "
            f"Chapter {chapter.get('chapterNumber','')} · "
            f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y')}"
        ),
    )

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "A concise revision sheet covering every key concept in this chapter. "
        "Use it for quick recap before quizzes."
    )
    ir.italic = True
    ir.font.color.rgb = MUTED

    for i, t in enumerate(chapter.get("extractedTopics", []) or [], start=1):
        # Topic heading
        h = doc.add_paragraph()
        hr = h.add_run(f"{i}. {t.get('title', 'Topic')}")
        hr.bold = True
        hr.font.size = Pt(14)
        hr.font.color.rgb = BRAND_BLUE

        # Meta line
        meta_bits = []
        if t.get("weight") is not None:
            meta_bits.append(f"Weight {int(float(t['weight']) * 100)}%")
        if meta_bits:
            m = doc.add_paragraph()
            mr = m.add_run(" · ".join(meta_bits))
            mr.font.size = Pt(9)
            mr.font.color.rgb = MUTED

        # Bulleted breakdown
        bullets = t.get("keyPoints") or _bullets_from_chunk(t.get("contentChunk", ""))
        for b in bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(b.rstrip("."))
        doc.add_paragraph()  # spacer

    # Footer
    f = doc.add_paragraph()
    fr = f.add_run(
        "— End of study sheet —  |  © Project Vidya · AI adaptive learning for CBSE"
    )
    fr.font.size = Pt(9)
    fr.italic = True
    fr.font.color.rgb = MUTED
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_question_paper_docx(chapter: dict, questions: list, include_answers: bool = True) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _header(
        doc,
        title=f"Practice Paper · {chapter.get('title','Chapter')}",
        subtitle=(
            f"Grade {chapter.get('grade','')} · {chapter.get('subject','')} · "
            f"Chapter {chapter.get('chapterNumber','')} · "
            f"{len(questions)} questions · "
            f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y')}"
        ),
    )

    # Group by topic
    topics_by_id = {t["topicId"]: t for t in chapter.get("extractedTopics", []) or []}
    by_topic: dict = {}
    for q in questions:
        by_topic.setdefault(q.get("topicId") or "_general", []).append(q)

    q_no = 1
    for tid, qs in by_topic.items():
        t = topics_by_id.get(tid)
        topic_title = t["title"] if t else "General"
        h = doc.add_paragraph()
        hr = h.add_run(f"Section: {topic_title}")
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = BRAND_BLUE

        for q in qs:
            # Question line
            p = doc.add_paragraph()
            qr = p.add_run(f"Q{q_no}. ")
            qr.bold = True
            p.add_run(q.get("questionText", ""))
            # Tags
            tags = []
            if q.get("type"): tags.append(q["type"])
            if q.get("bloomsTaxonomy"): tags.append(q["bloomsTaxonomy"])
            if q.get("difficultyLevel") is not None:
                tags.append(f"Difficulty {int(float(q['difficultyLevel'])*100)}%")
            if tags:
                tp = doc.add_paragraph()
                tr = tp.add_run(" · ".join(tags))
                tr.font.size = Pt(9)
                tr.font.color.rgb = MUTED

            # Options for MCQ
            if q.get("type") == "MCQ" and q.get("options"):
                for opt in q["options"]:
                    op = doc.add_paragraph(style="List Number")
                    op.add_run(opt.get("text", ""))

            q_no += 1
            doc.add_paragraph()

    # Answer key
    if include_answers:
        doc.add_page_break()
        ah = doc.add_paragraph()
        ahr = ah.add_run("Answer Key & Explanations")
        ahr.bold = True
        ahr.font.size = Pt(18)
        ahr.font.color.rgb = BRAND_BLUE
        doc.add_paragraph()

        q_no = 1
        for tid, qs in by_topic.items():
            for q in qs:
                p = doc.add_paragraph()
                qr = p.add_run(f"Q{q_no}. ")
                qr.bold = True

                if q.get("type") == "MCQ":
                    correct = q.get("correctOptionId")
                    correct_text = ""
                    for opt in (q.get("options") or []):
                        if opt.get("optionId") == correct:
                            correct_text = opt.get("text", "")
                            break
                    ans = p.add_run(f"Option {correct} — {correct_text}")
                    ans.bold = True
                    ans.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
                elif q.get("sampleAnswer"):
                    p.add_run("Sample Answer:").bold = True
                    doc.add_paragraph(q["sampleAnswer"])

                if q.get("explanation"):
                    e = doc.add_paragraph()
                    er = e.add_run("Explanation: ")
                    er.bold = True
                    er.font.color.rgb = MUTED
                    e.add_run(q["explanation"])

                q_no += 1
                doc.add_paragraph()

    # Footer
    f = doc.add_paragraph()
    fr = f.add_run(
        "— End of practice paper —  |  © Project Vidya · AI adaptive learning for CBSE"
    )
    fr.font.size = Pt(9)
    fr.italic = True
    fr.font.color.rgb = MUTED
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
